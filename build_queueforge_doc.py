from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "QueueForge_Interview_Deliverables.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run, size=size, color=color or RGBColor(0, 0, 0), bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)


def paragraph(text="", style=None, before=0, after=6, line=1.10, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def add_body(text, after=6, bold_first=None):
    p = paragraph(after=after)
    if bold_first and text.startswith(bold_first):
        r = p.add_run(bold_first)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_first):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run, size=10.8)
    return p


def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run, size=10.8)
    return p


def add_callout(title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_BLUE)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(body)
    set_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.6, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or value in ["Done", "Partial", "Recommended"] else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_architecture_diagram():
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [1.45, 1.75, 1.65, 1.65])
    rows = [
        ["React Dashboard", "Express API", "SQLite Store", "Worker Pool"],
        ["Auth, queue controls, job explorer, DLQ drawer", "JWT auth, validation, routes, cron sync", "Jobs, queues, retry policies, executions, logs, DLQ", "Heartbeat, atomic claim, execute, retry, failover"],
        ["User action", "POST /api/jobs", "QUEUED job", "CLAIMED then RUNNING"],
        ["Dashboard stats", "GET /api/dashboard/stats", "Aggregated health", "Active workers"],
        ["AI failure triage", "Gemini service", "DLQ summary stored", "Failure routed after max retries"],
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            fill = LIGHT_BLUE if r_idx == 0 else WHITE
            shade_cell(table.cell(r_idx, c_idx), fill)
            set_cell_text(table.cell(r_idx, c_idx), value, bold=r_idx == 0, size=9.4 if r_idx else 10.2, color=NAVY if r_idx == 0 else None, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_er_diagram():
    add_body("Entity relationship overview", after=4, bold_first="Entity relationship overview")
    rows = [
        ("users", "1 -> many", "projects"),
        ("projects", "1 -> many", "queues"),
        ("retry_policies", "1 -> many", "queues"),
        ("queues", "1 -> many", "jobs"),
        ("workers", "1 -> many", "jobs"),
        ("jobs", "1 -> many", "job_executions"),
        ("jobs", "1 -> many", "job_logs"),
        ("jobs", "1 -> 0/1", "dead_letter_queue"),
    ]
    add_matrix(["Parent Entity", "Relationship", "Child Entity"], rows, [2.0, 1.45, 3.05])


def setup_document():
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    section.header.paragraphs[0].text = "QueueForge Technical Submission"
    section.header.paragraphs[0].runs[0].font.size = Pt(9)
    section.header.paragraphs[0].runs[0].font.color.rgb = GRAY
    section.footer.paragraphs[0].text = "Prepared for internship interview review"
    section.footer.paragraphs[0].runs[0].font.size = Pt(9)
    section.footer.paragraphs[0].runs[0].font.color.rgb = GRAY


def build():
    setup_document()

    p = paragraph(before=8, after=4)
    r = p.add_run("QueueForge")
    set_font(r, size=28, color=NAVY, bold=True)
    p = paragraph(after=14)
    r = p.add_run("Distributed Job Scheduler - Interview Deliverables")
    set_font(r, size=14, color=GRAY)

    rows = [
        ("Candidate", "[Your Name]"),
        ("Project", "QueueForge - distributed job scheduling platform"),
        ("Stack", "Node.js, Express, SQLite, React, Vite, JWT, node-cron, Gemini API"),
        ("GitHub", "[Paste your GitHub repository link here]"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.35, 5.15])
    for label, value in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], LIGHT_GRAY)
        set_cell_text(cells[0], label, bold=True, color=NAVY)
        set_cell_text(cells[1], value)

    add_callout(
        "Executive summary",
        "QueueForge demonstrates a production-style background processing system with authenticated project isolation, queue-level concurrency, atomic worker claiming, retry backoff, dead-letter recovery, operational dashboards, and AI-assisted failure diagnosis. The implementation intentionally balances assignment scope with real distributed-systems trade-offs."
    )

    add_heading("Deliverables Coverage", 1)
    coverage = [
        ("Source code with setup instructions", "GitHub link placeholder included; README documents install, environment variables, API/worker startup, and demo flow."),
        ("Architecture diagram", "Included below as a layered runtime flow covering dashboard, API, database, cron engine, worker pool, DLQ, and AI service."),
        ("ER diagram", "Included below as a concise relationship matrix for users, projects, queues, policies, jobs, workers, executions, logs, and DLQ."),
        ("API documentation", "Included with endpoint purpose, auth requirement, core request fields, and response notes."),
        ("Design decisions and trade-offs", "Included with explicit notes on SQLite, atomic claiming, retry policies, cron polling, AI fallback, and production migration path."),
        ("Automated tests for critical functionality", "Included as a focused test plan because no test suite is currently present in the repository."),
    ]
    add_matrix(["Deliverable", "Where it is addressed"], coverage, [2.35, 4.15])

    add_heading("System Architecture", 1)
    add_body("QueueForge separates user interaction, API orchestration, persistence, scheduling, and job execution. The API owns validation and tenant boundaries; the worker process owns claim-and-execute behavior; SQLite provides durable assignment state for the demo environment.")
    add_architecture_diagram()

    add_heading("Runtime Flow", 2)
    for item in [
        "A user signs up or logs in and receives a JWT for authenticated project, queue, job, and dashboard operations.",
        "A project groups queues, while each queue has priority, pause/resume control, a concurrency limit, and a retry policy.",
        "A job is inserted as QUEUED, SCHEDULED, or delayed with run_at; scheduled jobs are instantiated by the cron engine.",
        "Workers send heartbeats, poll for eligible jobs, atomically flip jobs to CLAIMED, execute payloads, then mark results.",
        "Failures are retried with FIXED, LINEAR, or EXPONENTIAL backoff. Exhausted jobs move into the DLQ with captured payload, final error, logs, and AI summary.",
    ]:
        add_number(item)

    add_heading("Reliability and Concurrency", 2)
    add_bullet("Atomic claim protection: the worker uses BEGIN IMMEDIATE TRANSACTION and an UPDATE guarded by status='QUEUED', so racing workers cannot claim the same row.")
    add_bullet("Queue-level backpressure: each queue enforces its own concurrency_limit by counting CLAIMED/RUNNING jobs before claiming more work.")
    add_bullet("Priority scheduling: queues with higher priority are considered first, followed by due time and job id for deterministic ordering.")
    add_bullet("Graceful shutdown: workers stop polling, wait for active jobs to drain, and mark themselves DEAD before process exit.")
    add_bullet("Retry resilience: failed jobs are rescheduled with policy-driven backoff before final DLQ routing.")

    add_heading("Database Design", 1)
    add_er_diagram()
    add_body("The schema is deliberately normalized: user ownership flows through projects, queues carry retry policy and operational controls, jobs capture mutable lifecycle state, and append-only logs/executions preserve auditability. The DLQ isolates permanent failures from the active job path.")

    add_heading("Key Entities", 2)
    entities = [
        ("users", "Authenticated accounts with hashed passwords and created_at metadata."),
        ("projects", "User-owned containers that enforce tenant isolation for queues and jobs."),
        ("queues", "Execution lanes with priority, pause/resume, concurrency limits, and retry policy link."),
        ("retry_policies", "Reusable failure strategy: FIXED, LINEAR, or EXPONENTIAL with max retry and base delay."),
        ("jobs", "Core unit of work with type, status, payload, cron expression, run_at, retry_count, worker_id, and batch_id."),
        ("workers", "Runtime processors with heartbeat state used by dashboard health monitoring."),
        ("job_executions", "Attempt history for each job, including status, timestamps, worker, and error message."),
        ("job_logs", "Execution log stream used for observability and AI failure summarization."),
        ("dead_letter_queue", "Permanent failure sink with payload, last error, failed_at, and ai_failure_summary."),
    ]
    add_matrix(["Entity", "Responsibility"], entities, [1.75, 4.75])

    add_heading("API Documentation", 1)
    add_body("All endpoints except authentication require Authorization: Bearer <token>. Inputs are validated with express-validator and user ownership is enforced through project joins.")
    api_rows = [
        ("POST /api/auth/signup", "Create user account", "email, password", "201 with JWT token"),
        ("POST /api/auth/login", "Authenticate existing user", "email, password", "200 with JWT token"),
        ("POST /api/projects", "Create a project", "name", "201 project object"),
        ("GET /api/projects", "List user projects", "Bearer token", "Array of projects"),
        ("POST /api/queues", "Create queue with policy", "name, project_id, strategy, max_retries, base_delay_seconds, priority, concurrency_limit", "201 queue object"),
        ("GET /api/queues", "List queues with retry policy", "Bearer token", "Array of queues"),
        ("PATCH /api/queues/:id", "Update queue controls", "name, priority, concurrency_limit, paused", "Updated queue object"),
        ("POST /api/queues/:id/pause", "Pause dispatch", "queue id", "paused true"),
        ("POST /api/queues/:id/resume", "Resume dispatch", "queue id", "paused false"),
        ("POST /api/jobs", "Submit job", "queue_id, type, payload, cron_expression, delay_seconds, batch_id", "201 job id/status"),
        ("GET /api/jobs", "Paginated job list", "page, limit, status, queue_id", "page, limit, total, data"),
        ("GET /api/jobs/:id", "Job detail", "job id", "Job plus executions/logs/DLQ fields"),
        ("POST /api/jobs/:id/retry", "Manual retry", "FAILED job id", "Job requeued and DLQ row cleared"),
        ("GET /api/dashboard/stats", "Operational dashboard", "Bearer token", "Queue health, active workers, throughput, DLQ"),
    ]
    add_matrix(["Endpoint", "Purpose", "Primary Input", "Response"], api_rows, [1.65, 1.65, 1.95, 1.25])

    add_heading("Design Decisions and Trade-offs", 1)
    decisions = [
        ("SQLite for demo persistence", "Fast setup, zero external dependency, easy interview demo.", "Single-writer bottleneck; production should move to Postgres with row-level locks."),
        ("Atomic claim via guarded UPDATE", "Prevents duplicate execution when multiple workers race.", "SQLite locking is coarse-grained; Postgres SKIP LOCKED would scale better."),
        ("Queue-level concurrency", "Prevents one queue from overwhelming the worker pool.", "Current model counts active jobs at claim time; production could add worker capability routing."),
        ("Retry policies per queue", "Keeps failure handling configurable and easy to explain.", "Could evolve into per-job overrides and jittered backoff."),
        ("Cron sync every 30 seconds", "Simple, resilient scheduler loop with minimal moving parts.", "New schedule definitions may take up to 30 seconds to register."),
        ("AI summaries only on DLQ", "Keeps AI usage focused on high-value permanent failures.", "Requires Gemini key; code falls back gracefully when missing."),
        ("JWT authentication", "Simple stateless auth for API routes and frontend storage.", "Production should add refresh tokens, stronger secrets, and secure cookie options."),
    ]
    add_matrix(["Decision", "Why it helps", "Trade-off / next step"], decisions, [1.65, 2.25, 2.6])

    add_heading("Critical Test Plan", 1)
    add_body("The repository currently does not include an automated test suite, so the following test plan targets the assignment's highest-risk behavior. These are the tests I would implement first with Jest or Node's test runner plus a temporary SQLite database.")
    tests = [
        ("Atomic worker claim", "Seed one queued job, run two claim attempts concurrently, assert only one UPDATE returns changes=1 and one worker_id is persisted.", "Reliability & concurrency"),
        ("Retry backoff", "Fail a job under FIXED, LINEAR, and EXPONENTIAL policies; assert retry_count and run_at move correctly.", "Backend engineering"),
        ("DLQ transition", "Exhaust retries, assert job status FAILED, DLQ row created, logs captured, and AI fallback behaves without GEMINI_API_KEY.", "Critical failure path"),
        ("Queue pause/resume", "Pause queue, seed due job, assert worker does not claim; resume and assert job becomes claimable.", "Operational control"),
        ("Auth ownership", "Create two users and projects; assert one user cannot read or mutate the other's queues/jobs.", "Security boundary"),
        ("Dashboard aggregation", "Seed queue/job/execution/DLQ rows, assert /api/dashboard/stats returns correct counts and active worker heartbeat window.", "Frontend/API contract"),
        ("Manual retry", "Retry FAILED job, assert retry_count reset, status QUEUED, run_at current, and DLQ row removed.", "Recovery workflow"),
    ]
    add_matrix(["Test", "Scenario", "Evaluation Area"], tests, [1.75, 3.65, 1.1])

    add_heading("Evaluation Criteria Mapping", 1)
    criteria = [
        ("System Architecture", "20", "Layered API, worker, scheduler, dashboard, persistence, and AI failure analysis are clearly separated."),
        ("Database Design", "20", "Normalized schema with tenant ownership, queue policy, job lifecycle, execution history, logs, and DLQ."),
        ("Backend Engineering", "20", "JWT auth, validation, retry policy, atomic claims, graceful shutdown, and operational endpoints."),
        ("Reliability & Concurrency", "15", "Guarded UPDATE, queue concurrency limits, worker heartbeat, retries, DLQ, and shutdown drain."),
        ("Frontend & UX", "10", "React dashboard includes login, navigation, metrics, queues, jobs, workers, DLQ, and job detail drawer."),
        ("API Design", "5", "RESTful endpoint set with pagination, filtering, auth, ownership checks, and clear response shapes."),
        ("Documentation", "5", "README plus this deliverables document explain setup, flow, architecture, trade-offs, and tests."),
        ("Testing", "5", "Focused automated test plan identifies critical behavior to verify first."),
    ]
    add_matrix(["Evaluation Criteria", "Marks", "How QueueForge addresses it"], criteria, [1.65, 0.65, 4.2])

    add_heading("Setup and Demo Script", 1)
    add_number("Install dependencies with npm install at the repository root.")
    add_number("Create .env from .env.example and set JWT_SECRET plus GEMINI_API_KEY if AI summaries should call Gemini.")
    add_number("Run npm start for the API on port 3000.")
    add_number("Run npm run worker in one or more additional terminals to demonstrate concurrent worker claiming.")
    add_number("Use the documented API sequence: signup, create project, create queue, submit failing job, watch retries, inspect dashboard stats and DLQ summary.")

    add_heading("Recommended Production Evolution", 1)
    add_bullet("Replace SQLite with Postgres and use SELECT FOR UPDATE SKIP LOCKED for scalable multi-worker claims.")
    add_bullet("Add automated tests around worker claim races, retry scheduling, auth boundaries, DLQ, and dashboard aggregation.")
    add_bullet("Move secrets to a managed secret store and hard-fail when JWT_SECRET remains the development fallback in production.")
    add_bullet("Add job idempotency keys and payload schema validation for safer external integrations.")
    add_bullet("Add observability exports: structured logs, metrics counters, traces, and worker lag dashboards.")

    add_callout(
        "Interview positioning",
        "The strongest story to tell is not just that QueueForge runs jobs. It shows you understand the uncomfortable parts of backend systems: race conditions, retries, failure isolation, observability, tenant boundaries, and pragmatic trade-offs under assignment constraints."
    )

    doc.save(OUT)


doc = Document()
build()
print(OUT)
